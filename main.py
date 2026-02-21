"""
Render resume.yaml to HTML, PDF, and Word formats.
"""

import os
import yaml
from datetime import datetime
from pathlib import Path
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from weasyprint.text.fonts import FontConfiguration
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

source_file = os.getenv('SOURCE_FILE', 'resume.yaml')
output_html = os.getenv('OUTPUT_HTML', 'resume.html')
output_pdf = os.getenv('OUTPUT_PDF', 'resume.pdf')
output_docx = os.getenv('OUTPUT_DOCX', 'resume.docx')
output_md = os.getenv('OUTPUT_MD', 'readme.md')
templates_path = os.getenv('TEMPLATES_PATH', 'templates')


def format_date(date_str):
    """Format YYYY-MM-DD to Mon YYYY."""
    if not date_str:
        return ''
    try:
        return datetime.strptime(date_str[:7], '%Y-%m').strftime('%b %Y')
    except ValueError:
        return date_str[:4]


def load_resume(path):
    with open(path, encoding='utf-8') as f:
        return yaml.safe_load(f)


def make_env(templates_dir):
    env = Environment(loader=FileSystemLoader(templates_dir), keep_trailing_newline=True)
    env.filters['format_date'] = format_date
    return env


def render_template(env, template_name, data):
    return env.get_template(template_name).render(
        basics=data.get('basics', {}),
        work=data.get('work', []),
        education=data.get('education', []),
        certificates=data.get('certificates', []),
        skills=data.get('skills', []),
    )


def generate_pdf(html_string, output_path):
    font_config = FontConfiguration()
    HTML(string=html_string).write_pdf(output_path, font_config=font_config)


def generate_docx(data, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.875)
        section.right_margin = Inches(0.875)

    basics = data.get('basics', {})

    # Name
    heading = doc.add_heading(basics.get('name', ''), level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Label
    p = doc.add_paragraph(basics.get('label', ''))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Profiles
    profiles = basics.get('profiles', [])
    if profiles:
        contact = ' | '.join(f"{p['network']}: {p['url']}" for p in profiles)
        cp = doc.add_paragraph(contact)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    summary = basics.get('summary', '')
    if summary:
        doc.add_paragraph(summary.strip())

    # Experience
    doc.add_heading('Experience', level=1)
    for job in data.get('work', []):
        start = format_date(job.get('startDate', ''))
        end = format_date(job.get('endDate', '')) or 'Present'

        p = doc.add_paragraph()
        p.add_run(f"{job['name']} \u2014 {job.get('position', '')}").bold = True
        p.add_run(f"  |  {start} \u2013 {end}")

        for highlight in job.get('highlights', []):
            doc.add_paragraph(highlight, style='List Bullet')

    # Education
    doc.add_heading('Education', level=1)
    for edu in data.get('education', []):
        p = doc.add_paragraph()
        p.add_run(edu.get('institution', '')).bold = True
        end_year = edu.get('endDate', '')[:4]
        p.add_run(f" \u2014 {edu.get('studyType', '')}, {edu.get('area', '')} ({end_year})")

    # Certifications
    doc.add_heading('Certifications', level=1)
    for cert in data.get('certificates', []):
        p = doc.add_paragraph()
        p.add_run(cert.get('name', '')).bold = True
        year = cert.get('date', '')[:4]
        p.add_run(f", {cert.get('issuer', '')} ({year})")

    # Skills
    doc.add_heading('Skills', level=1)
    for skill in data.get('skills', []):
        p = doc.add_paragraph()
        p.add_run(f"{skill['name']}: ").bold = True
        p.add_run(', '.join(skill.get('keywords', [])))

    doc.save(output_path)


def main():
    data = load_resume(source_file)
    env = make_env(templates_path)

    md = render_template(env, 'resume.md.j2', data)
    Path(output_md).write_text(md, encoding='utf-8')
    print(f"Wrote {output_md}")

    html = render_template(env, 'resume.html.j2', data)
    Path(output_html).write_text(html, encoding='utf-8')
    print(f"Wrote {output_html}")

    generate_pdf(html, output_pdf)
    print(f"Wrote {output_pdf}")

    generate_docx(data, output_docx)
    print(f"Wrote {output_docx}")


if __name__ == '__main__':
    main()
